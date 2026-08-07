import os
import datetime
import requests
import docx
import base64
from google import genai
from google.genai import types

# 1. Récupération des clés d'API depuis les secrets GitHub
GEMINI_KEY = os.getenv("GEMINI_API_KEY")
BREVO_KEY = os.getenv("BREVO_API_KEY")

if not GEMINI_KEY or not BREVO_KEY:
    raise ValueError("Les clés API GEMINI_API_KEY ou BREVO_API_KEY ne sont pas définies dans les secrets GitHub.")

# Initialisation du client officiel Google GenAI
client = genai.Client(api_key=GEMINI_KEY)

# Prompt complet validé
PROMPT = """Rôle et objectif
Tu es un analyste spécialisé en politiques publiques et dynamiques de terrain du logement social et abordable. Tu produis une veille de presse quotidienne destinée à un expert du secteur. Appuie-toi sur Google Search pour ne citer que des sources réelles, récentes et vérifiables : n'invente jamais un article, une date ou un lien.

Périmètre géographique et thématique
- Pays : France, Allemagne, Italie.
- Sujet : Logement social et abordable au sens large (macro-politiques ET vie locale/terrain).
  • France : HLM, logement abordable/intermédiaire, bailleurs sociaux, financements (RLS, PLAI/PLUS/PLS, Action Logement, CDC), arbitrages maires/préfets/collectivités, décisions judiciaires/expulsions, tensions de terrain, vie des organismes, santé financière, associations de locataires, précarité énergétique.
  • Allemagne : sozialer Wohnungsbau, Sozialwohnungen, geförderter Wohnraum, Wohnraumförderung, kommunale/genossenschaftliche Wohnungsunternehmen, décisions des Länder et communes, initiatives syndicales (DGB/Mieterbund), reconversions, conflits d'usage.
  • Italie : edilizia residenziale pubblica (ERP), case popolari, housing sociale, canone calmierato, bandi régionaux/municipaux (ALER, ATER), syndicats de locataires (SUNIA, Unione Inquilini), expulsions (sfratti), réhabilitations et initiatives locales.

Sélection et signaux faibles
- Ne te limite PAS aux annonces ministérielles ou aux grandes lois. Recherche activement les SIGNAUX FAIBLES et la PRESSE RÉGIONALE/LOCALE : restructurations de bailleurs locaux, projets municipaux, arrêtés, conflits locataires/bailleurs, vagues d'expulsions, enjeux d'insalubrité/climat, prises de position d'acteurs de terrain.
- Sélectionne jusqu'à 20 informations au total, réparties selon l'actualité réelle, sans quota rigide par pays.
- Classe par ordre d'importance décroissante à l'intérieur de chaque pays.

Fenêtre temporelle stricte
- Uniquement les informations publiées dans les dernières 24 heures (horaire glissant jusqu'à maintenant).
- Écarte rigoureusement tout contenu publié il y a plus de 24 heures. Si une information majeure a plus de 24h et mérite d'être citée, inscris la mention obligatoire « [Antérieur à la période revue] ».
- Si la période est creuse sur un pays, indique-le clairement sans meubler.

Structure et mise en forme (Format compatible Word)
Structure le document de manière claire pour faciliter un copier-coller propre vers Microsoft Word :

# Actualité du Logement Social au [Date du jour]
**Heure de création :** [Heure actuelle] | **Périmètre :** France, Allemagne, Italie | **Informations retenues :** [Nombre total]

---

# 🇫🇷 France
[Sous-titre / Numérotation de l'item]
- **Titre :** [Titre clair et factuel]
- **Synthèse :** [Résumé de 2 à 3 phrases strictement factuel : chiffres, montants, noms d'organismes/acteurs, lieux précis, nature des faits ou des financements]
- **Source :** [Nom du média ou de l'institution], [Date exacte de publication] — [URL directe]

# 🇩🇪 Allemagne
[Même structure]

# 🇮🇹 Italie
[Même structure]

---
**Bilan de la veille :**
- Total d'items : [X]
- Pays sans actualité retenue dans la fenêtre des 24h : [Nom des pays le cas échéant]
"""

# 2. Appel à l'API Gemini avec modèle canonique et recherche Google
response = client.models.generate_content(
    model='gemini-1.5-pro',
    contents=PROMPT,
    config=types.GenerateContentConfig(
        tools=[types.Tool(google_search=types.GoogleSearch())]
    )
)

texte_veille = response.text if response.text else "Aucun contenu généré."

# 3. Génération du fichier Word (.docx)
doc = docx.Document()
doc.add_heading("Veille Logement Social", level=1)

for line in texte_veille.split('\n'):
    doc.add_paragraph(line)

today_str = datetime.date.today().strftime('%d/%m/%Y')
filename = f"Veille_Logement_Social_{datetime.date.today()}.docx"
doc.save(filename)

# 4. Envoi de l'email via Brevo
with open(filename, "rb") as f:
    encoded_file = base64.b64encode(f.read()).decode("utf-8")

brevo_url = "https://api.brevo.com/v3/smtp/email"
headers = {
    "accept": "application/json",
    "api-key": BREVO_KEY,
    "content-type": "application/json"
}

# ⚠️ N'oubliez pas de mettre votre email d'expéditeur validé dans Brevo !
EMAIL_SENDER = "ocfr@yahoo.fr" 

payload = {
    "sender": {"name": "Veille Logement", "email": EMAIL_SENDER},
    "to": [{"email": EMAIL_SENDER}],
    "subject": f"Actualité du Logement Social au {today_str}",
    "htmlContent": f"<h3>Bonjour,</h3><p>Voici la veille quotidienne sur le logement social du {today_str} en pièce jointe.</p>",
    "attachment": [
        {
            "content": encoded_file,
            "name": filename
        }
    ]
}

res = requests.post(brevo_url, json=payload, headers=headers)
print("Statut d'envoi Brevo :", res.status_code, res.text)
if res.status_code >= 400:
    raise Exception(f"Erreur d'envoi Brevo: {res.text}")
